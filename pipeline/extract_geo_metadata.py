from pathlib import Path
from docx import Document
import sys

sys.stdout.reconfigure(encoding="utf-8")

root = Path(__file__).resolve().parents[1] / "pipeline" / "geo"
for path in root.rglob("*.docx"):
    print(f"FILE={path}")
    doc = Document(path)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            print(text)
    for table in doc.tables:
        for row in table.rows:
            print(" | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells))
